# pylint: disable=too-many-lines

import io
from datetime import date
from datetime import datetime
from datetime import timezone
from http import HTTPStatus
from itertools import chain
from decimal import Decimal

import openpyxl
import pytest
import docx

from django.http.response import FileResponse
from django.urls import reverse
from openpyxl.utils import get_column_letter
from rest_framework.response import Response

from core.api.serializers.business_plan import BPActivityDetailSerializer
from core.api.tests.base import BaseTest
from core.api.tests.factories import AgencyFactory
from core.api.tests.factories import AlternativeTechnologyFactory
from core.api.tests.factories import BPActivityFactory
from core.api.tests.factories import FundingWindowFactory
from core.api.tests.factories import GroupFactory
from core.api.tests.factories import MetaProjectFactory
from core.api.tests.factories import MeetingFactory
from core.api.tests.factories import ProjectClusterFactory
from core.api.tests.factories import ProjectFactory
from core.api.tests.factories import ProjectSubSectorFactory
from core.api.tests.factories import ProjectTypeFactory
from core.api.tests.factories import SubstanceAltNameFactory
from core.api.tests.factories import SubstanceFactory
from core.api.export.single_project_v2.helpers import get_activity_data_from_instance
from core.api.export.single_project_v2.helpers import get_activity_data_from_json
from core.api.views import mya_export
from core.models import MetaProject
from core.models import Project
from core.models.business_plan import BPActivity
from core.models.project import ProjectOdsOdp
from core.models.project_metadata import ProjectField
from core.models.project_metadata import ProjectSpecificFields
from core.models.substance import Substance
from core.models.user import User

pytestmark = pytest.mark.django_db


def validate_docx_export(
    project: Project, user: User, response: FileResponse, content=None
):
    assert response.status_code == HTTPStatus.OK
    try:
        file_name = project.code.replace("/", "_")
    except AttributeError:
        file_name = f"project_{project.id}"
    assert response.filename == f"{file_name}.docx"
    assert (
        response.headers["Content-Type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    f = io.BytesIO(content if content is not None else response.getvalue())
    doc = docx.Document(f)
    assert f.tell() > 0

    to_find = [
        x
        for x in [
            user.get_full_name() or user.username,
            project.title,
            project.country.name,
            project.agency.name,
            project.cluster.name,
            project.total_fund,
        ]
        if x
    ]
    result = [False for v in to_find]

    paragraphs_to_check = chain(
        doc.sections[0].first_page_footer.paragraphs,
        doc.sections[0].even_page_footer.paragraphs,
        doc.sections[0].footer.paragraphs,
        doc.paragraphs,
    )

    found_description = False
    for t in doc.tables:
        if len(t.rows) and len(t.columns) == 1:
            if t.cell(0, 0).text == project.description:
                found_description = True

    assert (
        found_description
    ), f"Could not locate description ({project.description}) in output."

    for p in paragraphs_to_check:
        for i, v in enumerate(to_find):
            if result[i] is False:
                if v in p.text:
                    result[i] = True

    for t, r in zip(to_find, result):
        assert r is True, f"Could not locate {t} in output."


def get_docx_table_rows(content, header_text: str):
    f = io.BytesIO(content)
    doc = docx.Document(f)
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [cell.text for cell in table.rows[0].cells]
        if header_text in header_cells:
            return [[cell.text for cell in row.cells] for row in table.rows]
    raise AssertionError(f"Could not locate table with header {header_text}.")


def get_docx_field_values(content):
    doc = docx.Document(io.BytesIO(content))
    return {
        row.cells[0].text: row.cells[1].text
        for table in doc.tables
        for row in table.rows
        if len(row.cells) >= 2
    }


def ensure_substance_details_docx_fields(project: Project):
    project_specific_fields, _ = ProjectSpecificFields.objects.get_or_create(
        cluster=project.cluster,
        type=project.project_type,
        sector=project.sector,
    )
    substance_field = ProjectField.objects.create(
        import_name="ods_display_name",
        label="Substance",
        read_field_name="ods_display_name",
        write_field_name="ods_display_name",
        table="project_ods_odp",
        data_type="text",
        section="Substance Details",
    )
    project_specific_fields.fields.add(substance_field)


def ensure_obsolete_substance_details_docx_fields(project: Project):
    (
        project_specific_fields,
        _,
    ) = ProjectSpecificFields.objects.with_obsolete().get_or_create(
        cluster=project.cluster,
        type=project.project_type,
        sector=project.sector,
        defaults={"obsolete": True},
    )
    for field_name, label, data_type in [
        ("ods_display_name", "Substance", "text"),
        ("ods_replacement_text", "Replacement technologies", "text"),
        ("odp", "Phase-out (ODP tonnes)", "decimal"),
    ]:
        field = ProjectField.objects.create(
            import_name=field_name,
            label=label,
            read_field_name=field_name,
            write_field_name=field_name,
            table="ods_odp",
            data_type=data_type,
            section="Substance Details",
        )
        project_specific_fields.fields.add(field)


def validate_single_project_export(project: Project, response: FileResponse):
    assert response.filename == f"Project {project.id}.xlsx", response.filename

    wb = openpyxl.load_workbook(io.BytesIO(response.getvalue()))
    assert {
        "Identifiers",
        "Identifiers - BP Activity",
        "Cross-cutting",
        "Specific information - Overview",
        "Specific information - Substance details",
        "Impact",
    }.intersection(wb.sheetnames)
    sheet = wb["Identifiers"]
    assert sheet["A1"].value == "Country", sheet["A1"].value
    assert sheet["A2"].value == project.country.name, sheet["A2"].value

    sheet = wb["Identifiers - BP Activity"]
    assert sheet["A1"].value == "Activity ID", sheet["A1"].value
    assert sheet["A2"].value, sheet["A2"].value

    return wb


def validate_projects_export(project: Project, response: FileResponse):
    assert response.filename == "Projects.xlsx", response.filename

    wb = openpyxl.load_workbook(io.BytesIO(response.getvalue()))
    assert {
        "Projects",
        "Codes",
        "Legacy codes",
        "Metaproject codes",
        "Clusters",
        "Metaproject categories",
        "Project types",
        "Legacy project types",
        "Sectors",
        "Legacy sectors",
        "Subsectors",
        "Legacy subsectors",
        "Substance types",
        "Substances",
        "Status",
        "Serial numbers",
        "Legacy serial numbers",
        "Countries",
        "Titles",
    }.intersection(wb.sheetnames)
    sheet = wb["Projects"]
    assert sheet["A1"].value == "Code", sheet["A1"].value
    assert sheet["A2"].value == project.code, sheet["A2"].value


def get_inventory_headers(sheet):
    return {
        cell.value: get_column_letter(cell.column)
        for cell in sheet[1]
        if cell.value is not None
    }


def get_inventory_project_row(sheet, project_id):
    return next(
        (
            idx
            for idx in range(2, sheet.max_row + 1)
            if sheet[f"A{idx}"].value == project_id
        ),
        None,
    )


@pytest.fixture(name="project_with_linked_bp")
def _project_with_linked_bp(
    project: Project, bp_activity: BPActivity, substance: Substance
):
    project.bp_activity = bp_activity
    serialized = BPActivityDetailSerializer(bp_activity)
    ods_odp = {
        "ods_substance_id": substance.id,
        "odp": 11.11,
        "ods_replacement_text": "ods replacement test",
        "ods_replacement": None,
        "co2_mt": 323.23,
        "phase_out_mt": 123.23,
        "ods_type": "production",
        "sort_order": 1,
    }
    ProjectOdsOdp.objects.create(project=project, **ods_odp)
    project.bp_activity_json = serialized.data
    project.save()
    return project


@pytest.fixture(name="project_with_deleted_linked_bp")
def _project_with_deleted_linked_bp(project_with_linked_bp: Project):
    project_with_linked_bp.bp_activity.delete()
    project_with_linked_bp.bp_activity = None
    return project_with_linked_bp


class TestProjectV2ExportXLSX(BaseTest):
    url = reverse("project-v2-export")

    def test_bp_json_structure(self, project_with_linked_bp):
        # pylint: disable-next=import-outside-toplevel
        from core.api.serializers.project_v2 import ProjectDetailsV2Serializer

        serializer = ProjectDetailsV2Serializer(project_with_linked_bp)
        data = serializer.data
        activity_data_from_instance = get_activity_data_from_instance(data)
        activity_data_from_json = get_activity_data_from_json(data["bp_activity_json"])

        del activity_data_from_json["business_plan_id"]
        del activity_data_from_instance["business_plan_id"]

        assert activity_data_from_instance == activity_data_from_json

    def test_export_project_anon(self, project):
        self.client.force_authenticate(user=None)
        response: Response = self.client.get(self.url, {"project_id": project.id})
        assert response.status_code == HTTPStatus.FORBIDDEN, response.data

    def test_export_project_agency_submitter(
        self, project_with_linked_bp, agency_inputter_user
    ):
        self.client.force_authenticate(user=agency_inputter_user)
        response: FileResponse = self.client.get(
            self.url, {"project_id": project_with_linked_bp.id}
        )
        assert response.status_code == HTTPStatus.OK
        validate_single_project_export(project_with_linked_bp, response)

    def test_export_project_secretariat(
        self, project_with_linked_bp, secretariat_viewer_user, project_submitted_status
    ):
        self.client.force_authenticate(user=secretariat_viewer_user)
        project_with_linked_bp.submission_status = project_submitted_status
        project_with_linked_bp.save()
        response: FileResponse = self.client.get(
            self.url, {"project_id": project_with_linked_bp.id}
        )
        assert response.status_code == HTTPStatus.OK
        validate_single_project_export(project_with_linked_bp, response)

    def test_export_project_deleted_activity_secretariat(
        self,
        project_with_deleted_linked_bp,
        secretariat_viewer_user,
        project_submitted_status,
    ):
        self.client.force_authenticate(user=secretariat_viewer_user)
        project_with_deleted_linked_bp.submission_status = project_submitted_status
        project_with_deleted_linked_bp.save()
        response: FileResponse = self.client.get(
            self.url, {"project_id": project_with_deleted_linked_bp.id}
        )

        assert response.status_code == HTTPStatus.OK
        assert (
            project_with_deleted_linked_bp.bp_activity is None
        ), project_with_deleted_linked_bp.bp_activity
        assert (
            project_with_deleted_linked_bp.bp_activity_json
        ), project_with_deleted_linked_bp.bp_activity_json
        validate_single_project_export(project_with_deleted_linked_bp, response)

    def test_export_projects_anon(self, project):
        self.client.force_authenticate(user=None)
        response: Response = self.client.get(self.url, {"project_id": project.id})
        assert response.status_code == HTTPStatus.FORBIDDEN, response.data

    def test_export_projects_agency_submitter(
        self, project, agency_inputter_user, project_approved_status
    ):
        # Set submission status to approved to have the project code shown in the export
        project.submission_status = project_approved_status
        project.save()
        self.client.force_authenticate(user=agency_inputter_user)
        response: FileResponse = self.client.get(self.url)
        assert response.status_code == HTTPStatus.OK
        validate_projects_export(project, response)

    def test_export_projects_secretariat(
        self, project, secretariat_viewer_user, project_approved_status
    ):
        # Set submission status to approved to have the project code shown in the export
        project.submission_status = project_approved_status
        project.save()
        self.client.force_authenticate(user=secretariat_viewer_user)
        response: FileResponse = self.client.get(self.url)
        assert response.status_code == HTTPStatus.OK
        validate_projects_export(project, response)

    def test_export_inventory_report_reads_project_objects_directly(
        self, admin_user, project_approved_status
    ):
        lead_agency = AgencyFactory.create(name="Lead agency")
        cluster = ProjectClusterFactory.create(name="Cluster A")
        funding_window = FundingWindowFactory.create(description="Window A")
        meta_project = MetaProjectFactory.create(
            type=MetaProject.MetaProjectType.MYA,
            end_date=datetime(2026, 12, 31, 0, 0, 0, tzinfo=timezone.utc),
        )
        subsector_one = ProjectSubSectorFactory.create(name="Subsector A")
        subsector_two = ProjectSubSectorFactory.create(name="Subsector B")
        project = ProjectFactory.create(
            version=3,
            metacode="META-123",
            code="PRJ-123",
            legacy_code="LEG-123",
            title="Inventory project",
            description="Inventory description",
            excom_provision="Provision text",
            products_manufactured="Manufactured product",
            tranche=4,
            additional_funding=True,
            production=True,
            lead_agency=lead_agency,
            funding_window=funding_window,
            meta_project=meta_project,
            cluster=cluster,
            subsectors=[subsector_one, subsector_two],
            submission_status=project_approved_status,
        )
        bp_activity = BPActivityFactory.create(
            agency=project.agency,
            country=project.country,
            initial_id=123,
            lvc_status=BPActivity.LVCStatus.undefined,
            status=BPActivity.Status.undefined,
        )
        project.bp_activity = bp_activity
        project.sector_legacy = "Sector legacy"
        project.subsector_legacy = "Subsector legacy"
        project.save()

        self.client.force_authenticate(user=admin_user)
        response: FileResponse = self.client.get(self.url, {"inventory_report": "true"})

        assert response.status_code == HTTPStatus.OK

        wb = openpyxl.load_workbook(io.BytesIO(response.getvalue()))
        sheet = wb["Projects"]
        headers = get_inventory_headers(sheet)
        row = get_inventory_project_row(sheet, project.id)
        last_header = sheet.cell(1, sheet.max_column).value

        assert row is not None
        assert sheet[f"{headers['id']}{row}"].value == project.id
        assert sheet[f"{headers['Country']}{row}"].value == project.country.name
        assert sheet[f"{headers['Meta Code']}{row}"].value == project.metacode
        assert sheet[f"{headers['Code']}{row}"].value == project.code
        assert sheet[f"{headers['Legacy Code']}{row}"].value == project.legacy_code
        assert sheet[f"{headers['Agency']}{row}"].value == project.agency.name
        assert sheet[f"{headers['Lead Agency']}{row}"].value == lead_agency.name
        assert sheet[f"{headers['Cluster']}{row}"].value == project.cluster.name
        assert sheet[f"{headers['Type']}{row}"].value == project.project_type.code
        assert sheet[f"{headers['Sector']}{row}"].value == project.sector.name
        assert sheet[f"{headers['Sector Legacy']}{row}"].value == project.sector_legacy
        assert sheet[f"{headers['Subsector']}{row}"].value == ", ".join(
            [subsector_one.name, subsector_two.name]
        )
        assert (
            sheet[f"{headers['Subsector Legacy']}{row}"].value
            == project.subsector_legacy
        )
        assert sheet[f"{headers['Project Title']}{row}"].value == project.title
        assert (
            sheet[f"{headers['Project Description']}{row}"].value == project.description
        )
        assert (
            sheet[f"{headers['ExCom Provision']}{row}"].value == project.excom_provision
        )
        assert (
            sheet[f"{headers['Product manufactured']}{row}"].value
            == project.products_manufactured
        )
        assert sheet[f"{headers['Tranche number']}{row}"].value == project.tranche
        assert sheet[f"{headers['Category']}{row}"].value == "MYA"
        assert sheet[f"{headers['Funding window']}{row}"].value == str(
            funding_window.decision.number
        )
        assert sheet[f"{headers['Production']}{row}"].value == "Yes"
        assert (
            sheet[f"{headers['Business plan activity']}{row}"].value
            == bp_activity.get_display_internal_id
        )
        assert sheet[f"{headers['Additional funding']}{row}"].value == "Yes"
        assert last_header == "cost_effectiveness_co2_actual"
        assert (
            sheet[f"{headers['MYA Completion Date']}{row}"].value.date()
            == meta_project.end_date.date()
        )
        assert sheet[f"{headers['Transfer Meeting']}{row}"].value in (None, "")
        assert sheet[f"{headers['Transfer Decision']}{row}"].value in (None, "")
        assert sheet[f"{headers['Agency Transferred From']}{row}"].value in (None, "")

    def test_export_inventory_report_uses_ods_replacement_name(
        self, admin_user, project_approved_status
    ):
        project = ProjectFactory.create(
            version=3,
            submission_status=project_approved_status,
        )
        substance = SubstanceFactory.create(name="HCFC-22")
        replacement = AlternativeTechnologyFactory.create(name="Methyl formate")
        ProjectOdsOdp.objects.create(
            project=project,
            ods_substance=substance,
            ods_display_name="",
            ods_replacement_text="",
            ods_replacement=replacement,
            odp=3.4,
            sort_order=1,
        )

        self.client.force_authenticate(user=admin_user)
        response: FileResponse = self.client.get(self.url, {"inventory_report": "true"})

        assert response.status_code == HTTPStatus.OK

        wb = openpyxl.load_workbook(io.BytesIO(response.getvalue()))
        sheet = wb["Projects"]
        headers = get_inventory_headers(sheet)
        row = get_inventory_project_row(sheet, project.id)

        assert row is not None
        assert sheet[f"{headers['ODS Name 1']}{row}"].value == "HCFC-22"
        assert sheet[f"{headers['ODS Replacement 1']}{row}"].value == "Methyl formate"

    def test_export_inventory_report_only_includes_approved_projects(
        self, admin_user, project_approved_status, project_recommended_status
    ):
        approved_project = ProjectFactory.create(
            version=3,
            code="APPROVED-CODE",
            submission_status=project_approved_status,
        )
        recommended_project = ProjectFactory.create(
            version=3,
            code="RECOMMENDED-CODE",
            submission_status=project_recommended_status,
        )

        self.client.force_authenticate(user=admin_user)
        response: FileResponse = self.client.get(self.url, {"inventory_report": "true"})

        assert response.status_code == HTTPStatus.OK

        wb = openpyxl.load_workbook(io.BytesIO(response.getvalue()))
        sheet = wb["Projects"]

        assert get_inventory_project_row(sheet, approved_project.id) is not None
        assert get_inventory_project_row(sheet, recommended_project.id) is None

    def test_export_inventory_report_uses_mya_type_revised_date_for_extended_date(
        self, admin_user, project_approved_status
    ):
        meta_project = MetaProjectFactory.create(
            type=MetaProject.MetaProjectType.MYA,
            end_date=datetime(2025, 12, 1, tzinfo=timezone.utc),
        )
        inv_type = ProjectTypeFactory.create(code="INV")
        tas_type = ProjectTypeFactory.create(code="TAS")
        inv_project = ProjectFactory.create(
            version=3,
            category=Project.Category.MYA,
            meta_project=meta_project,
            project_type=inv_type,
            project_end_date=date(2021, 12, 1),
            submission_status=project_approved_status,
        )
        tas_project = ProjectFactory.create(
            version=3,
            category=Project.Category.MYA,
            meta_project=meta_project,
            project_type=tas_type,
            project_end_date=date(2018, 12, 1),
            submission_status=project_approved_status,
        )
        ProjectFactory.create(
            version=3,
            category=Project.Category.MYA,
            meta_project=meta_project,
            project_type=tas_type,
            date_comp_revised=date(2020, 12, 1),
            project_end_date=date(2020, 12, 1),
            submission_status=project_approved_status,
        )

        self.client.force_authenticate(user=admin_user)
        response: FileResponse = self.client.get(self.url, {"inventory_report": "true"})

        assert response.status_code == HTTPStatus.OK

        wb = openpyxl.load_workbook(io.BytesIO(response.getvalue()))
        sheet = wb["Projects"]
        headers = get_inventory_headers(sheet)
        inv_row = get_inventory_project_row(sheet, inv_project.id)
        tas_row = get_inventory_project_row(sheet, tas_project.id)

        assert inv_row is not None
        assert tas_row is not None
        assert sheet[f"{headers['Extended date']}{inv_row}"].value.date() == date(
            2025, 12, 1
        )
        assert sheet[f"{headers['Extended date']}{tas_row}"].value.date() == date(
            2020, 12, 1
        )

    @pytest.mark.parametrize(
        "stored_end_date, project_end_date, expected_end_date",
        [
            (None, date(2026, 12, 1), date(2026, 12, 1)),
            (
                datetime(2025, 6, 30, tzinfo=timezone.utc),
                date(2030, 12, 1),
                date(2025, 6, 30),
            ),
        ],
    )
    def test_export_inventory_report_derives_mya_completion_date(
        self,
        admin_user,
        project_approved_status,
        stored_end_date,
        project_end_date,
        expected_end_date,
    ):
        meta_project = MetaProjectFactory.create(
            type=MetaProject.MetaProjectType.MYA,
            end_date=stored_end_date,
        )
        project = ProjectFactory.create(
            version=3,
            category=Project.Category.MYA,
            meta_project=meta_project,
            project_end_date=project_end_date,
            submission_status=project_approved_status,
        )

        self.client.force_authenticate(user=admin_user)
        response: FileResponse = self.client.get(self.url, {"inventory_report": "true"})

        assert response.status_code == HTTPStatus.OK

        wb = openpyxl.load_workbook(io.BytesIO(response.getvalue()))
        sheet = wb["Projects"]
        headers = get_inventory_headers(sheet)
        row = get_inventory_project_row(sheet, project.id)

        assert row is not None
        assert (
            sheet[f"{headers['MYA Completion Date']}{row}"].value.date()
            == expected_end_date
        )

    def test_export_inventory_report_derives_substance_from_ods_and_legacy_sector(
        self, admin_user, project_approved_status
    ):
        hcfc_group = GroupFactory.create(group_id="CI")
        hcfc_substance = SubstanceFactory.create(name="HCFC-22", group=hcfc_group)
        hcfc_project = ProjectFactory.create(
            version=3,
            substance_type="CFC",
            submission_status=project_approved_status,
        )
        ProjectOdsOdp.objects.create(
            project=hcfc_project,
            ods_substance=hcfc_substance,
        )

        ctc_group = GroupFactory.create(group_id="BII")
        ctc_substance = SubstanceFactory.create(
            name="Carbon Tetrachloride test", group=ctc_group
        )
        ctc_project = ProjectFactory.create(
            version=3,
            substance_type="CFC",
            submission_status=project_approved_status,
        )
        ctc_display_name = f"CTC-{ctc_project.id}"
        SubstanceAltNameFactory.create(name=ctc_display_name, substance=ctc_substance)
        ProjectOdsOdp.objects.create(
            project=ctc_project,
            ods_display_name=ctc_display_name,
        )
        halon_project = ProjectFactory.create(
            version=3,
            substance_type="CFC",
            sector_legacy="HAL",
            submission_status=project_approved_status,
        )
        methyl_bromide_project = ProjectFactory.create(
            version=3,
            substance_type="CFC",
            sector_legacy="FUM",
            submission_status=project_approved_status,
        )

        self.client.force_authenticate(user=admin_user)
        response: FileResponse = self.client.get(self.url, {"inventory_report": "true"})

        assert response.status_code == HTTPStatus.OK

        wb = openpyxl.load_workbook(io.BytesIO(response.getvalue()))
        sheet = wb["Projects"]
        headers = get_inventory_headers(sheet)

        expected = {
            hcfc_project.id: "HCFC",
            ctc_project.id: "CTC",
            halon_project.id: "Halon",
            methyl_bromide_project.id: "Methyl Bromide",
        }
        for project_id, substance in expected.items():
            row = get_inventory_project_row(sheet, project_id)
            assert row is not None
            assert sheet[f"{headers['Substance']}{row}"].value == substance

    def test_export_inventory_report_populates_prior_meeting_columns(
        self, admin_user, project_approved_status
    ):
        final_project = ProjectFactory.create(
            version=6,
            code="PRIOR-MEETING-CODE",
            total_fund=160,
            support_cost_psc=16,
            meeting=MeetingFactory.create(number=306),
            post_excom_meeting=MeetingFactory.create(number=206),
            date_approved=date(2024, 6, 15),
            submission_status=project_approved_status,
        )
        ProjectFactory.create(
            latest_project=final_project,
            version=3,
            total_fund=100,
            support_cost_psc=10,
            meeting=MeetingFactory.create(number=303),
            date_approved=date(2024, 3, 15),
        )
        ProjectFactory.create(
            latest_project=final_project,
            version=4,
            total_fund=120,
            support_cost_psc=12,
            meeting=MeetingFactory.create(number=304),
            post_excom_meeting=MeetingFactory.create(number=204),
            date_approved=date(2024, 4, 15),
        )
        ProjectFactory.create(
            latest_project=final_project,
            version=5,
            total_fund=140,
            support_cost_psc=14,
            meeting=MeetingFactory.create(number=305),
            post_excom_meeting=MeetingFactory.create(number=205),
            date_approved=date(2024, 5, 15),
        )

        self.client.force_authenticate(user=admin_user)
        response: FileResponse = self.client.get(self.url, {"inventory_report": "true"})

        assert response.status_code == HTTPStatus.OK

        wb = openpyxl.load_workbook(io.BytesIO(response.getvalue()))
        sheet = wb["Projects"]
        headers = get_inventory_headers(sheet)
        row = get_inventory_project_row(sheet, final_project.id)

        assert row is not None

        expected = {
            1: {"fund": 100, "psc": 10, "meeting": 303, "date": date(2024, 3, 15)},
            2: {"fund": 20, "psc": 2, "meeting": 204, "date": date(2024, 4, 15)},
            3: {"fund": 20, "psc": 2, "meeting": 205, "date": date(2024, 5, 15)},
        }
        currency_format = "#,##0;-#,##0;;@"
        for idx, values in expected.items():
            fund_cell = sheet[f"{headers[f'Funds Approved {idx}']}{row}"]
            psc_cell = sheet[f"{headers[f'Support Costs Approved {idx}']}{row}"]

            assert fund_cell.value == values["fund"]
            assert fund_cell.number_format == currency_format
            assert psc_cell.value == values["psc"]
            assert psc_cell.number_format == currency_format
            assert (
                sheet[f"{headers[f'Meeting Approved {idx}']}{row}"].value
                == values["meeting"]
            )
            assert (
                sheet[f"{headers[f'Date Approved {idx}']}{row}"].value.date()
                == values["date"]
            )

    def test_export_inventory_report_populates_adjustment_columns(
        self, admin_user, project_approved_status
    ):
        final_project = ProjectFactory.create(
            version=6,
            code="ADJUSTMENT-CODE",
            total_fund=160,
            support_cost_psc=16,
            meeting=MeetingFactory.create(number=306),
            post_excom_meeting=MeetingFactory.create(number=206),
            date_actual=date(2024, 6, 15),
            date_approved=date(2024, 6, 15),
            adjustment=True,
            fund_transferred=9,
            psc_transferred=3,
            submission_status=project_approved_status,
        )
        ProjectFactory.create(
            latest_project=final_project,
            version=3,
            total_fund=100,
            support_cost_psc=10,
            meeting=MeetingFactory.create(number=303),
            date_actual=date(2024, 3, 15),
        )
        ProjectFactory.create(
            latest_project=final_project,
            version=4,
            total_fund=120,
            support_cost_psc=12,
            meeting=MeetingFactory.create(number=304),
            post_excom_meeting=MeetingFactory.create(number=204),
            date_actual=date(2024, 4, 15),
            date_approved=date(2024, 4, 15),
            adjustment=True,
            fund_transferred=7,
            psc_transferred=1,
        )
        ProjectFactory.create(
            latest_project=final_project,
            version=5,
            total_fund=140,
            support_cost_psc=14,
            meeting=MeetingFactory.create(number=305),
            post_excom_meeting=MeetingFactory.create(number=205),
            date_actual=date(2024, 5, 15),
            date_approved=date(2024, 5, 15),
            adjustment=True,
            fund_transferred=8,
            psc_transferred=2,
        )

        self.client.force_authenticate(user=admin_user)
        response: FileResponse = self.client.get(self.url, {"inventory_report": "true"})

        assert response.status_code == HTTPStatus.OK

        wb = openpyxl.load_workbook(io.BytesIO(response.getvalue()))
        sheet = wb["Projects"]
        headers = get_inventory_headers(sheet)
        row = get_inventory_project_row(sheet, final_project.id)

        assert row is not None

        expected = {
            1: {"fund": 20, "psc": 2, "meeting": 204, "date": date(2024, 4, 15)},
            2: {"fund": 20, "psc": 2, "meeting": 205, "date": date(2024, 5, 15)},
            3: {"fund": 20, "psc": 2, "meeting": 206, "date": date(2024, 6, 15)},
        }
        currency_format = "#,##0;-#,##0;;@"
        for idx, values in expected.items():
            fund_cell = sheet[f"{headers[f'Fund Adjustments {idx}']}{row}"]
            psc_cell = sheet[f"{headers[f'Support Cost Adjustments {idx}']}{row}"]

            assert fund_cell.value == values["fund"]
            assert fund_cell.number_format == currency_format
            assert psc_cell.value == values["psc"]
            assert psc_cell.number_format == currency_format
            assert (
                sheet[f"{headers[f'Adjustments Meeting {idx}']}{row}"].value
                == values["meeting"]
            )
            assert (
                sheet[f"{headers[f'Adjustments Date {idx}']}{row}"].value.date()
                == values["date"]
            )

    def test_export_mya_adds_filters_and_totals(
        self,
        secretariat_viewer_user,
        project_approved_status,
    ):
        first_meta_project = MetaProjectFactory.create(
            type=MetaProject.MetaProjectType.MYA
        )
        first_project = ProjectFactory.create(
            meta_project=first_meta_project,
            category=Project.Category.MYA,
            submission_status=project_approved_status,
        )
        first_meta_project.umbrella_code = "META-001"
        first_meta_project.country = first_project.country
        first_meta_project.cluster = first_project.cluster
        first_meta_project.project_duration = 12
        first_meta_project.project_funding = Decimal("100.50")
        first_meta_project.support_cost = Decimal("10.25")
        first_meta_project.project_cost = Decimal("110.75")
        first_meta_project.target_odp = Decimal("2.50")
        first_meta_project.baseline_odp = Decimal("3.75")
        first_meta_project.number_of_smes_directly_funded = 4
        first_meta_project.cost_effectiveness_kg = Decimal("5.50")
        first_meta_project.save()

        second_meta_project = MetaProjectFactory.create(
            type=MetaProject.MetaProjectType.MYA,
            umbrella_code="META-002",
            project_duration=24,
            project_funding=Decimal("200.25"),
            support_cost=Decimal("20.50"),
            project_cost=Decimal("220.75"),
            target_odp=Decimal("3.50"),
            baseline_odp=Decimal("4.25"),
            number_of_smes_directly_funded=6,
            cost_effectiveness_kg=Decimal("6.50"),
        )
        second_project = ProjectFactory.create(
            meta_project=second_meta_project,
            category=Project.Category.MYA,
            submission_status=project_approved_status,
        )
        second_meta_project.country = second_project.country
        second_meta_project.cluster = second_project.cluster
        second_meta_project.save()

        self.client.force_authenticate(user=secretariat_viewer_user)
        response: FileResponse = self.client.get(
            self.url, {"category": [Project.Category.MYA]}
        )

        assert response.status_code == HTTPStatus.OK

        wb = openpyxl.load_workbook(io.BytesIO(response.getvalue()))
        sheet = wb["MYAs"]
        assert sheet.auto_filter.ref == "A1:AA3"
        assert sheet["A4"].value is None, "Spacer row not empty."
        assert sheet["A5"].value == "Subtotal (filtered results)"
        assert sheet["A6"].value == "Grand total"

        assert sheet["D5"].value is None
        assert sheet["D6"].value is None

        for idx, header in enumerate(mya_export.HEADERS, start=1):
            if header.get("in_grand_total"):
                ltr = get_column_letter(idx)
                assert sheet[f"{ltr}5"].value == f"=SUBTOTAL(9,{ltr}2:{ltr}3)"
                assert sheet[f"{ltr}6"].value == f"=SUM({ltr}2:{ltr}3)"

                if cell_format := header.get("cell_format"):
                    assert sheet[f"{ltr}5"].number_format == cell_format
                    assert sheet[f"{ltr}6"].number_format == cell_format

    def test_export_mya_uses_computed_values_only_when_stored_value_missing(
        self,
        secretariat_viewer_user,
        project_approved_status,
        project_draft_status,
    ):
        fallback_meta_project = MetaProjectFactory.create(
            type=MetaProject.MetaProjectType.MYA,
            umbrella_code="META-FALLBACK",
        )
        fallback_first_project = ProjectFactory.create(
            meta_project=fallback_meta_project,
            category=Project.Category.MYA,
            submission_status=project_approved_status,
            total_fund=100,
            support_cost_psc=10,
            project_start_date=date(2024, 1, 15),
            project_end_date=date(2024, 6, 15),
        )
        fallback_second_project = ProjectFactory.create(
            meta_project=fallback_meta_project,
            category=Project.Category.MYA,
            submission_status=project_approved_status,
            total_fund=150,
            support_cost_psc=15,
            project_start_date=date(2024, 2, 15),
            project_end_date=date(2024, 7, 20),
        )
        fallback_meta_project.country = fallback_first_project.country
        fallback_meta_project.cluster = fallback_first_project.cluster
        fallback_meta_project.save()
        ProjectOdsOdp.objects.create(
            project=fallback_first_project,
            co2_mt=20,
            odp=2,
            phase_out_mt=3,
        )
        ProjectOdsOdp.objects.create(
            project=fallback_first_project,
            co2_mt=7,
            odp=1,
            phase_out_mt=2,
        )
        ProjectOdsOdp.objects.create(
            project=fallback_second_project,
            co2_mt=30,
            odp=4,
            phase_out_mt=5,
        )
        draft_project = ProjectFactory.create(
            meta_project=fallback_meta_project,
            category=Project.Category.MYA,
            submission_status=project_draft_status,
            total_fund=500,
            support_cost_psc=50,
            project_start_date=date(2024, 1, 1),
            project_end_date=date(2024, 12, 31),
        )
        ProjectOdsOdp.objects.create(
            project=draft_project,
            co2_mt=100,
            odp=10,
            phase_out_mt=11,
        )
        archived_project = ProjectFactory.create(
            meta_project=fallback_meta_project,
            latest_project=fallback_first_project,
            category=Project.Category.MYA,
            submission_status=project_approved_status,
            total_fund=999,
            support_cost_psc=99,
        )
        ProjectOdsOdp.objects.create(
            project=archived_project,
            co2_mt=50,
            odp=5,
            phase_out_mt=6,
        )

        override_meta_project = MetaProjectFactory.create(
            type=MetaProject.MetaProjectType.MYA,
            umbrella_code="META-OVERRIDE",
            start_date=datetime(2023, 5, 1, tzinfo=timezone.utc),
            end_date=datetime(2023, 9, 1, tzinfo=timezone.utc),
            project_funding=Decimal("999.50"),
            support_cost=Decimal("88.25"),
            phase_out_co2_eq_t=Decimal("77.75"),
            phase_out_odp=Decimal("6.50"),
            phase_out_mt=Decimal("5.25"),
        )
        override_project = ProjectFactory.create(
            meta_project=override_meta_project,
            category=Project.Category.MYA,
            submission_status=project_approved_status,
            total_fund=200,
            support_cost_psc=20,
            project_start_date=date(2024, 3, 1),
            project_end_date=date(2024, 10, 1),
        )
        override_meta_project.country = override_project.country
        override_meta_project.cluster = override_project.cluster
        override_meta_project.save()
        ProjectOdsOdp.objects.create(
            project=override_project,
            co2_mt=40,
            odp=8,
            phase_out_mt=9,
        )

        self.client.force_authenticate(user=secretariat_viewer_user)
        response: FileResponse = self.client.get(
            self.url, {"category": [Project.Category.MYA]}
        )

        assert response.status_code == HTTPStatus.OK

        wb = openpyxl.load_workbook(io.BytesIO(response.getvalue()))
        sheet = wb["MYAs"]
        headers = get_inventory_headers(sheet)

        fallback_row = next(
            idx
            for idx in range(2, sheet.max_row + 1)
            if sheet[f"{headers['Metacode']}{idx}"].value == "META-FALLBACK"
        )
        assert (
            sheet[f"{headers['Start date (MYA)']}{fallback_row}"].value == "15/01/2024"
        )
        assert sheet[f"{headers['End date (MYA)']}{fallback_row}"].value == "20/07/2024"
        assert (
            sheet[
                f"{headers['MYA Total agreed funding in principle (US $)']}{fallback_row}"
            ].value
            == 250
        )
        assert (
            sheet[
                f"{headers['MYA Total support costs in principle (US $)']}{fallback_row}"
            ].value
            == 25
        )
        assert (
            sheet[f"{headers['Phase-out (CO2-eq tonnes) (MYA)']}{fallback_row}"].value
            == 57
        )
        assert (
            sheet[f"{headers['Phase-out (ODP tonnes) (MYA)']}{fallback_row}"].value == 7
        )
        assert (
            sheet[f"{headers['Phase-out (metric tonnes) (MYA)']}{fallback_row}"].value
            == 10
        )

        override_row = next(
            idx
            for idx in range(2, sheet.max_row + 1)
            if sheet[f"{headers['Metacode']}{idx}"].value == "META-OVERRIDE"
        )
        assert (
            sheet[f"{headers['Start date (MYA)']}{override_row}"].value == "01/05/2023"
        )
        assert sheet[f"{headers['End date (MYA)']}{override_row}"].value == "01/09/2023"
        assert (
            sheet[
                f"{headers['MYA Total agreed funding in principle (US $)']}{override_row}"
            ].value
            == 999.5
        )
        assert (
            sheet[
                f"{headers['MYA Total support costs in principle (US $)']}{override_row}"
            ].value
            == 88.25
        )
        assert (
            sheet[f"{headers['Phase-out (CO2-eq tonnes) (MYA)']}{override_row}"].value
            == 77.75
        )
        assert (
            sheet[f"{headers['Phase-out (ODP tonnes) (MYA)']}{override_row}"].value
            == 6.5
        )
        assert (
            sheet[f"{headers['Phase-out (metric tonnes) (MYA)']}{override_row}"].value
            == 5.25
        )

    def test_export_mya_uses_first_approved_project_lead_agency(
        self,
        secretariat_viewer_user,
        project_approved_status,
    ):
        first_lead_agency = AgencyFactory.create(name="Lead agency A")
        second_lead_agency = AgencyFactory.create(name="Lead agency B")
        meta_project = MetaProjectFactory.create(
            type=MetaProject.MetaProjectType.MYA,
            umbrella_code="META-LEAD",
        )
        first_project = ProjectFactory.create(
            meta_project=meta_project,
            category=Project.Category.MYA,
            submission_status=project_approved_status,
            lead_agency=first_lead_agency,
        )
        ProjectFactory.create(
            meta_project=meta_project,
            category=Project.Category.MYA,
            submission_status=project_approved_status,
            lead_agency=second_lead_agency,
        )
        meta_project.country = first_project.country
        meta_project.cluster = first_project.cluster
        meta_project.save()

        self.client.force_authenticate(user=secretariat_viewer_user)
        response: FileResponse = self.client.get(
            self.url, {"category": [Project.Category.MYA]}
        )

        assert response.status_code == HTTPStatus.OK

        wb = openpyxl.load_workbook(io.BytesIO(response.getvalue()))
        sheet = wb["MYAs"]
        headers = get_inventory_headers(sheet)
        row = next(
            idx
            for idx in range(2, sheet.max_row + 1)
            if sheet[f"{headers['Metacode']}{idx}"].value == "META-LEAD"
        )

        assert sheet[f"{headers['Lead agency']}{row}"].value == first_lead_agency.name


class TestProjectV2ExportDOCX(BaseTest):
    url = reverse("project-v2-export")

    def test_export_project_anon(self, project):
        self.client.force_authenticate(user=None)
        response: Response = self.client.get(self.url, {"project_id": project.id})
        assert response.status_code == HTTPStatus.FORBIDDEN, response.data

    def test_export_project_agency_submitter(
        self, project_with_linked_bp, agency_inputter_user
    ):
        ensure_substance_details_docx_fields(project_with_linked_bp)
        project_with_linked_bp.products_manufactured = "Manufactured product"
        project_with_linked_bp.save()
        project_ods_odp = project_with_linked_bp.ods_odp.first()
        project_ods_odp.ods_display_name = "Text baseline substance"
        project_ods_odp.ods_replacement_text = "Text replacement substance"
        project_ods_odp.save()
        fallback_replacement = AlternativeTechnologyFactory.create(
            name="FK replacement substance"
        )
        ProjectOdsOdp.objects.create(
            project=project_with_linked_bp,
            ods_display_name="FK baseline substance",
            ods_replacement_text="",
            ods_replacement=fallback_replacement,
        )

        self.client.force_authenticate(user=agency_inputter_user)
        response: FileResponse = self.client.get(
            self.url, {"project_id": project_with_linked_bp.id, "output_format": "docx"}
        )
        content = response.getvalue()
        validate_docx_export(
            project_with_linked_bp, agency_inputter_user, response, content
        )

        products_manufactured_rows = get_docx_table_rows(
            content, "Products manufactured"
        )
        assert products_manufactured_rows[0] == [
            "Products manufactured",
            "Manufactured product",
        ]

        substance_rows = get_docx_table_rows(content, "Baseline substance")
        substance_columns = [row[:2] for row in substance_rows]
        assert [
            "Text baseline substance",
            "Text replacement substance",
        ] in substance_columns
        assert [
            "FK baseline substance",
            "FK replacement substance",
        ] in substance_columns

        project_with_linked_bp.products_manufactured = ""
        project_with_linked_bp.save()
        response = self.client.get(
            self.url, {"project_id": project_with_linked_bp.id, "output_format": "docx"}
        )
        content = response.getvalue()
        validate_docx_export(
            project_with_linked_bp, agency_inputter_user, response, content
        )

        with pytest.raises(AssertionError):
            get_docx_table_rows(content, "Products manufactured")

        substance_rows = get_docx_table_rows(content, "Baseline substance")
        substance_columns = [row[:2] for row in substance_rows]
        assert [
            "Text baseline substance",
            "Text replacement substance",
        ] in substance_columns
        assert [
            "FK baseline substance",
            "FK replacement substance",
        ] in substance_columns

    def test_export_project_docx_uses_obsolete_specific_fields_for_ods_rows(
        self, project_with_linked_bp, agency_inputter_user
    ):
        ensure_obsolete_substance_details_docx_fields(project_with_linked_bp)
        project_ods_odp = project_with_linked_bp.ods_odp.first()
        project_ods_odp.ods_display_name = "CFC-11"
        project_ods_odp.ods_replacement_text = "Cyclopentane"
        project_ods_odp.odp = 200
        project_ods_odp.save()

        self.client.force_authenticate(user=agency_inputter_user)
        response: FileResponse = self.client.get(
            self.url, {"project_id": project_with_linked_bp.id, "output_format": "docx"}
        )

        content = response.getvalue()
        validate_docx_export(
            project_with_linked_bp, agency_inputter_user, response, content
        )
        substance_rows = get_docx_table_rows(content, "Baseline substance")
        assert any(
            row[0] == "CFC-11" and row[1] == "Cyclopentane" and row[-1] == "200.00"
            for row in substance_rows
        )

    def test_export_project_mya_uses_computed_values_only_when_stored_value_missing(
        self,
        secretariat_viewer_user,
        project_approved_status,
        project_draft_status,
    ):
        cluster = ProjectClusterFactory.create(category=Project.Category.MYA)
        meta_project = MetaProjectFactory.create(
            type=MetaProject.MetaProjectType.MYA,
            cluster=cluster,
            project_funding=Decimal("999.50"),
            phase_out_odp=Decimal("6.50"),
        )
        project_specs = [
            (
                project_approved_status,
                100,
                10,
                date(2024, 1, 15),
                date(2024, 6, 15),
                20,
                2,
                3,
            ),
            (
                project_approved_status,
                150,
                15,
                date(2024, 2, 15),
                date(2024, 7, 20),
                30,
                4,
                5,
            ),
            (
                project_draft_status,
                500,
                50,
                date(2024, 1, 1),
                date(2024, 12, 31),
                100,
                10,
                11,
            ),
        ]
        projects = []
        for (
            status,
            fund,
            support,
            start,
            end,
            co2_mt,
            odp,
            phase_out_mt,
        ) in project_specs:
            mya_project = ProjectFactory.create(
                meta_project=meta_project,
                category=Project.Category.MYA,
                cluster=cluster,
                submission_status=status,
                total_fund=fund,
                support_cost_psc=support,
                project_start_date=start,
                project_end_date=end,
            )
            ProjectOdsOdp.objects.create(
                project=mya_project,
                co2_mt=co2_mt,
                odp=odp,
                phase_out_mt=phase_out_mt,
            )
            projects.append(mya_project)

        self.client.force_authenticate(user=secretariat_viewer_user)
        response: FileResponse = self.client.get(
            self.url, {"project_id": projects[0].id, "output_format": "docx"}
        )

        assert response.status_code == HTTPStatus.OK
        field_values = get_docx_field_values(response.getvalue())
        assert field_values["Start date (MYA)"] == "15/01/2024"
        assert field_values["End date (MYA)"] == "20/07/2024"
        assert field_values["MYA Total agreed funding in principle (US $)"] == "$999.50"
        assert field_values["MYA Total support costs in principle (US $)"] == "$25.00"
        assert Decimal(field_values["Phase-out (CO2-eq tonnes) (MYA)"]) == Decimal("50")
        assert Decimal(field_values["Phase-out (ODP tonnes) (MYA)"]) == Decimal("6.5")
        assert Decimal(field_values["Phase-out (metric tonnes) (MYA)"]) == Decimal("8")
