# pylint:disable=protected-access

import datetime
import io
from itertools import chain

import docx
from django.http import FileResponse
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.oxml import CT_P
from docx.oxml import CT_Tbl
from docx.table import Table

from core.api.export import TEMPLATE_DIR
from core.api.export.single_project_v2.docx_headers import get_headers_metaproject
from core.api.export.single_project_v2.helpers import format_decimal
from core.api.export.single_project_v2.xlsx_headers import get_headers_cross_cutting
from core.api.export.single_project_v2.xlsx_headers import (
    get_headers_specific_information,
)
from core.api.serializers.meta_project import MetaProjecMyaDetailsSerializer
from core.api.serializers.project_v2 import ProjectDetailsV2Serializer
from core.models import Project
from core.models import ProjectSpecificFields, ProjectField
from core.models import User


def document_response(doc, filename):
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    res = FileResponse(out, as_attachment=True, filename=filename)
    return res


class ProjectsV2ProjectExportDocx:
    user = User
    project: Project
    doc: docx.Document
    template_path = (
        TEMPLATE_DIR
        / "Word Template for data entered into the system during project submission online.docx"
    )

    _known_tables: list[tuple[CT_P, Table]]

    def __init__(self, project, user):
        self.user = user
        self.project = project
        self._known_tables = []
        with self.template_path.open("rb") as tpl:
            self.doc = docx.Document(tpl)

    def find_table(self, after_p_text="", exclude_from_cleanup=False):
        found = None
        found_p = None

        after_p_text = after_p_text.strip().lower()

        if after_p_text:
            for e in self.doc.element.body:
                if isinstance(e, CT_P) and after_p_text == e.text.strip().lower():
                    found_p = e
                elif isinstance(e, CT_Tbl) and found_p:
                    found = Table(e, self.doc)
                    break

        if found and found_p and not exclude_from_cleanup:
            self._known_tables.append((found_p, found))

        return found

    @staticmethod
    def remove_empty_table(table: Table, paragraph: CT_P):
        table._element.getparent().remove(table._element)
        paragraph.getparent().remove(paragraph)

    def remove_empty_tables(self):
        for p, table in self._known_tables:
            if len(table.rows) == 1:
                self.remove_empty_table(table, p)

    def clean_paragraphs(self, keep=3):
        consecutive_empty = 0
        to_remove = []
        for p in self.doc.paragraphs:
            if not p.text.strip():
                consecutive_empty += 1
                if consecutive_empty > keep:
                    to_remove.append(p)
            else:
                consecutive_empty = 0

        for p in to_remove:
            p._element.getparent().remove(p._element)

    def build_front_page(self, data):

        # check footers
        for section in self.doc.sections:
            footer_paragraphs = chain(
                section.first_page_footer.paragraphs,
                section.even_page_footer.paragraphs,
                section.footer.paragraphs,
            )
            for p in footer_paragraphs:
                if p.text.startswith("Generated on"):
                    now = datetime.datetime.now(datetime.UTC).strftime("%d/%m/%Y")
                    user = self.user.get_full_name() or self.user.username
                    agency = self.user.agency.name if self.user.agency else ""
                    p.text = f"Generated on {now} by {user}" + (
                        f'of "{agency}"' if agency else ""
                    )
                    p.runs[0].italic = True

        # check paragraphs
        for p in self.doc.paragraphs:
            p_style = {
                "italic": False,
                "bold": False,
                "underline": False,
            }

            if p.runs:
                for k in p_style:
                    p_style[k] = getattr(p.runs[0], k, False)

            if p.text.startswith("Project Title"):
                p.text = self.project.title
            elif p.text.startswith("Country:"):
                p.add_run(data.get("country", "-"), None)
            elif p.text.startswith("Agency:"):
                p.add_run(data.get("agency", "-"), None)
            elif p.text.startswith("Lead Agency:"):
                p.add_run(data.get("lead_agency", "-"), None)
            elif p.text.startswith("Type:"):
                p.add_run(data.get("project_type", {}).get("name", "-"), None)
            elif p.text.startswith("Sector:"):
                p.add_run(data.get("sector", {}).get("name", "-"), None)
            elif p.text.startswith("Cluster:"):
                p.add_run(data.get("cluster", {}).get("name", "-"), None)
            elif p.text.startswith("Project costs:"):
                p.add_run(format_decimal(data.get("total_fund", "-")), None)
            elif p.text.startswith("Support costs:"):
                p.add_run(format_decimal(data.get("support_cost_psc", "-")), None)
            elif p.text.startswith("Metacode:"):
                p.add_run(data.get("metacode", "-"), None)
            elif p.text.startswith("Code:"):
                p.add_run(data.get("code", "-"), None)

            if p.runs:
                for k, v in p_style.items():
                    setattr(p.runs[0], k, v)

        description_table = self.find_table(
            "Project Description:", exclude_from_cleanup=True
        )
        if description_table:
            description_table.cell(0, 0).text = data.get("description", "")

    def _write_header_to_table(self, headers, table, data):
        for header in headers:
            row = table.add_row()
            for c_idx, cell in enumerate(row.cells):
                is_boolean = header.get("type") == "bool"
                is_decimal = header.get("type") == "decimal"
                is_dollar_value = (
                    header.get("type") == "number" or is_decimal
                ) and header.get("cell_format")

                value = data.get(header["id"])
                if c_idx == 0:
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(header["headerName"])
                    if header.get("docx_highlight", False):
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                elif c_idx == 1 and header.get("method"):
                    cell.text = header["method"](data, header)
                elif c_idx == 1 and is_dollar_value:
                    cell.text = (
                        format_decimal(
                            value, with_decimals=header.get("docx_decimals", True)
                        )
                        or ""
                    )
                elif c_idx == 1 and is_decimal:
                    cell.text = format_decimal(value, is_currency=False) or ""
                elif c_idx == 1 and is_boolean:
                    cell.text = "Yes" if value else "No"
                elif c_idx == 1:
                    cell.text = str(value or "")

    def _write_substance_table(self, _, table, data):
        substances = data.get("ods_odp", [])
        data_fields = [
            "ods_display_name",
            "???",
            "ods_replacement",
        ]

        for field, label in [
            ("phase_out_mt", "Phase out MT"),
            ("co2_mt", "Phase out CO2"),
            ("odp", "Phase out ODP"),
        ]:
            has_data = any(s[field] is not None for s in substances)

            if has_data:
                data_fields.append(field)

                table.add_column(70)
                cell = table.rows[0].cells[-1]
                cell.text = ""
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run(label)
                run.bold = True
                run.underline = True

        for substance in substances:
            row = table.add_row()
            for c_idx, cell in enumerate(row.cells):
                field = data_fields[c_idx]
                value = substance.get(field, "")
                if field in ["phase_out_mt", "co2_mt", "odp"]:
                    value = format_decimal(value, is_currency=False)

                cell.text = str(value or "")

    def _write_impact_target_actual(self, project, headers, table, data):
        planned_headers = {}
        actual_headers = {}
        for header in headers:
            if header["id"].endswith("_actual"):
                planned_name = header["id"].split("_actual")[0]
                actual_headers[planned_name] = header
            else:
                planned_headers[header["id"]] = header

        for field_id, header in planned_headers.items():
            actual_id = actual_headers.get(field_id, {}).get("id")
            row = table.add_row()
            row_data = [
                project.code,
                header["headerName"],
                data.get(field_id, ""),
                data.get(actual_id) if actual_id else "",
            ]
            for c_idx, cell in enumerate(row.cells):
                cell.text = str(row_data[c_idx] or "")

    def build_related_tranches(self):
        table = self.find_table("Related tranches (metacode and linked projects)")
        is_mya = self.project.category == Project.Category.MYA
        if table and is_mya and self.project.meta_project:
            related_projects = Project.objects.filter(
                meta_project__id=self.project.meta_project.id,
            )
            row = table.add_row()
            for project in related_projects:
                row_data = [
                    project.metacode,
                    project.code,
                    project.status.name,
                ]
                for c_idx, cell in enumerate(row.cells):
                    cell.text = str(row_data[c_idx] or "")

    def _write_project_cross_cutting_fields(
        self,
        table=None,
        fields=None,
        data=None,
        writer=None,
    ):
        writer = self._write_header_to_table if not writer else writer
        if data and table and fields:
            headers = get_headers_cross_cutting(fields, for_docx=True)
            writer(headers, table, data)

    def build_cross_cutting(self, data):
        self._write_project_cross_cutting_fields(
            table=self.find_table("Cross-cutting fields"),
            fields=self._get_fields_for_section(section_name="Cross-Cutting").exclude(
                label="Description"
            ),
            data=data,
        )

    def _get_fields_for_section(
        self,
        section_name: str,
        fields_obj: ProjectSpecificFields | None = None,
        **filters,
    ):
        if not fields_obj:
            return ProjectField.objects.get_visible_fields_for_user(self.user).filter(
                section__in=[section_name], **filters
            )
        return fields_obj.fields.get_visible_fields_for_user(self.user).filter(
            section__in=[section_name], **filters
        )

    def _write_project_specific_fields(
        self,
        table=None,
        fields=None,
        data=None,
        writer=None,
    ):
        writer = self._write_header_to_table if not writer else writer
        if data and table and fields:
            headers = get_headers_specific_information(fields)
            writer(headers, table, data)

    def _write_metaproject_fields(
        self,
        table=None,
        data=None,
    ):
        writer = self._write_header_to_table
        if data and table:
            headers = get_headers_metaproject()
            writer(headers, table, data)

    def build_specific_information(self, data):
        project_specific_fields_obj = ProjectSpecificFields.objects.filter(
            cluster=self.project.cluster,
            type=self.project.project_type,
            sector=self.project.sector,
        ).first()

        if project_specific_fields_obj:
            self._write_project_specific_fields(
                table=self.find_table("Project specific fields"),
                fields=self._get_fields_for_section(
                    "Header", project_specific_fields_obj
                ),
                data=data,
            )
            self._write_project_specific_fields(
                table=self.find_table("Substance details"),
                fields=self._get_fields_for_section(
                    "Substance Details", project_specific_fields_obj
                ),
                data=data,
                writer=self._write_substance_table,
            )
            self._write_project_specific_fields(
                table=self.find_table("Impact indicators"),
                fields=self._get_fields_for_section(
                    "Impact",
                    project_specific_fields_obj,
                    is_actual=False,
                ),
                data=data,
            )

    def build_impact_previous_tranches(self):
        table = self.find_table("Impact (previous MYA tranches) If applicable")
        if table and self.project.tranche and self.project.meta_project:
            related_projects = Project.objects.filter(
                country=self.project.country,
                cluster=self.project.cluster,
                tranche__lt=self.project.tranche,
                submission_status__name="Approved",
            )
            for project in related_projects:
                project_specific_fields_obj = ProjectSpecificFields.objects.filter(
                    cluster=project.cluster,
                    type=project.project_type,
                    sector=project.sector,
                ).first()

                fields = self._get_fields_for_section(
                    "Impact",
                    project_specific_fields_obj,
                )
                data = ProjectDetailsV2Serializer(project).data

                headers = get_headers_specific_information(fields)
                self._write_impact_target_actual(project, headers, table, data)

    def build_mya(self):
        metaproject = self.project.meta_project
        metaproject_data = {}

        if metaproject:
            metaproject_data = MetaProjecMyaDetailsSerializer(metaproject).data

        self._write_metaproject_fields(
            table=self.find_table("MYA"),
            data=metaproject_data,
        )

    def remove_page_breaks(self):
        # Iterate through all paragraphs
        for paragraph in self.doc.paragraphs:
            # Check each run in the paragraph
            for run in paragraph.runs:
                # Access the XML element
                for elem in run._element:
                    # Check if it's a page break (w:br with w:type="page")
                    if elem.tag.endswith("br"):
                        if (
                            elem.get(
                                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                            )
                            == "page"
                        ):
                            # Remove the page break element
                            elem.getparent().remove(elem)

    def build_document(self):
        serializer = ProjectDetailsV2Serializer(self.project)
        data = serializer.data

        self.build_front_page(data)
        self.build_related_tranches()
        self.build_cross_cutting(data)
        self.build_specific_information(data)
        self.build_impact_previous_tranches()
        self.build_mya()

        self.remove_page_breaks()
        self.remove_empty_tables()
        self.clean_paragraphs(keep=3)

    def export_docx(self):
        self.build_document()
        try:
            filename = self.project.code.replace("/", "_")
        except AttributeError:
            filename = f"project_{self.project.id}"
        return document_response(self.doc, filename=f"{filename}.docx")
